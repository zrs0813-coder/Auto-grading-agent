#!/usr/bin/env python3
"""
Automated grading agent for MIA 5100 group project reports.

Workflow
--------
1. Read and understand the Project Instruction document.
2. Derive a track-specific rubric (30 points) for each of the two tracks,
   once per run, cached to rubrics.json so every group in a track is graded
   against an identical rubric.
3. For each group folder: load the report (PDF or DOCX), classify its track,
   then grade it criterion by criterion against that track's rubric.
4. Emit a structured CSV, an Excel workbook, and a readable Markdown report.

LLM calls go through the Geotab GenAI Gateway, which is OpenAI-API compatible:
point the openai SDK's base_url at https://genai-{region}.geotab.com/api/v2 and
pass the gateway token as api_key (sent as Authorization: Bearer).
"""

import argparse
import csv
import json
import os
import re
import time
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional, Tuple

import docx
import openai
import pandas as pd
import PyPDF2
from dotenv import load_dotenv
from openpyxl.utils import get_column_letter

# ============================================================================
# CONFIGURATION
# ============================================================================

# Gateway credentials live in .env (gitignored) - see .env for the fields.
load_dotenv(Path(__file__).resolve().parent / ".env")

API_KEY = os.environ.get("GENAI_API_KEY", "")
BASE_URL = os.environ.get("GENAI_BASE_URL", "https://genai-us.geotab.com/api/v2")
# Fallback only - .env normally supplies GENAI_MODEL. Keep this in sync with
# .env so a missing/absent .env does not silently grade with a weaker model.
MODEL = os.environ.get("GENAI_MODEL", "claude-opus-5")

# Root folder holding the 12 group subfolders.
ROOT_FOLDER = (
    "MIA 5100 Project Report Download Aug 8, 2026 1247 PM (Group Submission Folder)/"
    "MIA 5100 Project Report Download Aug 8, 2026 1247 PM (Group Submission Folder)"
)

# The Project Instruction document (used to derive both rubrics).
INSTRUCTIONS_FILE = "MIA5100 - Project - Instructions.pdf"

# Total points available per report.
TOTAL_POINTS = 30

# The report must follow the IEEE conference paper template. This criterion is
# fixed in code rather than generated, so it is always present and always
# carries exactly this weight. The remaining points are derived from the
# instruction document.
IEEE_POINTS = 2
IEEE_CRITERION = {
    "criterion": "IEEE Conference Format & Structure",
    "max_points": IEEE_POINTS,
    "description": (
        "The report is presented as an IEEE conference paper. Judge this on "
        "structural markers only: an abstract at the top, index terms/keywords, "
        "numbered sections and subsections, figure/table captions, and a numbered "
        "IEEE-style reference list. Award FULL credit if the report is broadly "
        "IEEE-like. A single missing element (for example index terms) is NOT a "
        "deduction. Deduct only if the report plainly disregards the template "
        "altogether - for example an unstructured essay with no sections and no "
        "formal reference list."
    ),
}

TEMPERATURE = 0.0
SEED = 42
TIMEOUT = 300
MAX_RETRIES = 3

# Token ceilings per call type. Criterion-level grading with evidence quotes
# needs considerably more room than a plain score.
MAX_TOKENS_RUBRIC = 6000
MAX_TOKENS_CLASSIFY = 500
MAX_TOKENS_GRADE = 12000

# Characters of report text sent to the model. Reports are ~7 IEEE pages;
# this is generous headroom while bounding cost on pathological submissions.
MAX_REPORT_CHARS = 120000

REPORT_EXTENSIONS = (".pdf", ".docx", ".doc")

TRACKS = {
    1: "Hands-On AI System Development Project (Build & Evaluate)",
    2: "AI System Design Project (Analyze & Design)",
}

# The grading policy is deliberately lenient. This text is injected into both
# the rubric-generation and grading prompts so the two stay consistent.
GRADING_POLICY = """
GRADING PHILOSOPHY - READ CAREFULLY AND FOLLOW EXACTLY:

This grading must be VERY LENIENT and GENEROUS. The default outcome for any
criterion is FULL CREDIT.

- Give FULL credit whenever the student reasonably satisfies the requirement.
- DO NOT deduct for grammar, spelling, style, organization, presentation, or
  page count anywhere.
- Document format is assessed ONLY within the single dedicated
  "IEEE Conference Format & Structure" criterion, and even there leniently.
  No other criterion may lose points for formatting reasons.
- DO NOT deduct for reasonable alternative approaches. If the team solved the
  problem a different but defensible way, that earns full credit.
- DO NOT deduct for depth you merely wish were greater. "Could be more
  detailed" is NOT a valid deduction.
- Only deduct when there is a CLEAR and SUBSTANTIVE violation of the rubric
  criterion - for example, a required component is entirely absent, or what is
  present is fundamentally incorrect.
- If something is slightly incomplete but the core requirement is satisfied,
  prefer NO deduction, or at most a small one (1 point).
- DO NOT go looking for reasons to deduct points. Do not nitpick.
- EVERY deduction must be justified by specific evidence quoted or cited from
  the submission. If you cannot point to concrete evidence of a substantive
  failure, you MUST give full credit.
- If there is no substantive issue, give full credit.
- When genuinely uncertain, give full credit.
"""


# ============================================================================
# GATEWAY CLIENT
# ============================================================================


class GatewayClient:
    """Thin wrapper over the OpenAI SDK pointed at the Geotab GenAI Gateway."""

    def __init__(self, api_key: str, base_url: str = BASE_URL, model: str = MODEL):
        self.model = model
        self.base_url = base_url
        self.client = openai.OpenAI(
            api_key=api_key,
            base_url=base_url,  # <-- routes through the Geotab GenAI Gateway
            timeout=TIMEOUT,
            max_retries=0,  # retries handled in call() below
        )

    def call(self, system_prompt: str, user_prompt: str, max_tokens: int) -> str:
        """Send a chat completion to the gateway and return the assistant text."""
        params: Dict[str, Any] = {
            "model": self.model,
            "messages": [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            "temperature": TEMPERATURE,
            "max_tokens": max_tokens,
        }

        # `seed` is OpenAI-specific; Claude and Gemini reject it. temperature=0
        # is what actually drives consistency for those models.
        if self.model.startswith("gpt-"):
            params["seed"] = SEED

        last_error = None
        for attempt in range(MAX_RETRIES):
            try:
                response = self.client.chat.completions.create(**params)
                choice = response.choices[0]
                if choice.finish_reason == "length":
                    print(
                        f"  [gateway] WARNING: response hit the {max_tokens}-token cap "
                        "and was truncated"
                    )
                return (choice.message.content or "").strip()
            except Exception as e:
                last_error = e
                message = str(e)

                # Some upstream models reject seed/temperature. Drop the
                # offending param and retry immediately.
                dropped = False
                if "400" in message or "unsupported" in message.lower():
                    for param in ("seed", "temperature", "max_tokens"):
                        if param in params and param in message:
                            params.pop(param)
                            print(f"  [gateway] dropping unsupported param '{param}', retrying")
                            dropped = True
                            break
                if dropped:
                    continue

                if attempt < MAX_RETRIES - 1:
                    backoff = 2**attempt
                    print(f"  [gateway] {message} - retrying in {backoff}s")
                    time.sleep(backoff)

        raise RuntimeError(f"GenAI Gateway request failed: {last_error}")


# ============================================================================
# DOCUMENT LOADING
# ============================================================================


def extract_pdf(path: Path) -> str:
    with open(path, "rb") as f:
        reader = PyPDF2.PdfReader(f)
        return "\n".join((page.extract_text() or "") for page in reader.pages)


def extract_docx(path: Path) -> str:
    document = docx.Document(str(path))
    parts = [p.text for p in document.paragraphs]
    # Tables carry results/comparison content that paragraphs alone would miss.
    for table in document.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells]
            if any(cells):
                parts.append(" | ".join(cells))
    return "\n".join(parts)


def extract_text(path: Path) -> str:
    """Extract text from a PDF or DOCX report. Returns '' on failure."""
    try:
        suffix = path.suffix.lower()
        if suffix == ".pdf":
            return extract_pdf(path)
        if suffix in (".docx", ".doc"):
            return extract_docx(path)
        print(f"  Unsupported file type: {path.name}")
        return ""
    except Exception as e:
        print(f"  Error reading {path.name}: {e}")
        return ""


def group_sort_key(folder: Path) -> Tuple[int, str]:
    """Sort group folders numerically (Group 2 before Group 10)."""
    match = re.search(r"Group\s*(\d+)", folder.name, re.IGNORECASE)
    return (int(match.group(1)) if match else 9999, folder.name)


def group_label(folder: Path) -> str:
    match = re.search(r"Group\s*(\d+)", folder.name, re.IGNORECASE)
    return f"Group {int(match.group(1))}" if match else folder.name


def find_report(folder: Path) -> Optional[Path]:
    """Find the single report file inside a group folder (recursively)."""
    candidates = [
        p
        for p in sorted(folder.rglob("*"))
        if p.is_file()
        and p.suffix.lower() in REPORT_EXTENSIONS
        and not p.name.startswith("~$")  # Word lock files
    ]
    if not candidates:
        return None
    # Prefer a file that looks like the report if several are present.
    for p in candidates:
        if re.search(r"report|final", p.name, re.IGNORECASE):
            return p
    return candidates[0]


def discover_groups(root: Path) -> List[Tuple[str, Path, Optional[Path]]]:
    """Return [(group_label, folder, report_path_or_None)] for each subfolder."""
    folders = sorted((p for p in root.iterdir() if p.is_dir()), key=group_sort_key)
    return [(group_label(f), f, find_report(f)) for f in folders]


# ============================================================================
# JSON PARSING
# ============================================================================


def parse_json(text: str) -> Dict[str, Any]:
    """Parse a model response, tolerating markdown fences and stray prose."""
    text = text.strip()
    fenced = re.search(r"```(?:json)?\s*(.*?)\s*```", text, re.DOTALL)
    if fenced:
        text = fenced.group(1).strip()
    try:
        return json.loads(text)
    except json.JSONDecodeError:
        start, end = text.find("{"), text.rfind("}")
        if start != -1 and end > start:
            return json.loads(text[start : end + 1])
        raise


# ============================================================================
# GRADING AGENT
# ============================================================================


class ProjectGrader:
    def __init__(self, client: GatewayClient, instructions_text: str):
        self.client = client
        self.instructions = instructions_text
        self.rubrics: Dict[int, Dict[str, Any]] = {}

    # --- step 1+2: understand instructions, build track rubrics -------------

    def build_rubric(self, track: int) -> Dict[str, Any]:
        """Derive a TOTAL_POINTS rubric for one track from the instructions.

        The IEEE format criterion is fixed in code; the model generates only the
        remaining content criteria, so the format weight can never drift.
        """
        other = 2 if track == 1 else 1
        content_points = TOTAL_POINTS - IEEE_POINTS
        system = (
            "You are a university course designer. You read assignment instructions "
            "and turn them into precise grading rubrics. Always respond with valid JSON only."
        )
        user = f"""
Below is the full Project Instruction document for MIA 5100 Machine Learning.

The project has TWO tracks (streams):
  Track 1: {TRACKS[1]}
  Track 2: {TRACKS[2]}

Build the CONTENT grading criteria for the REPORT deliverable of
**Track {track}: {TRACKS[track]}**.

Requirements for the criteria you produce:
- They must apply to Track {track} ONLY. Do not include criteria that belong to Track {other}.
- Derive them from what the instruction document actually requires of the
  report for this track (including the shared sections required of both tracks).
- Use between 6 and 8 criteria.
- The max_points across your criteria MUST sum to EXACTLY {content_points}.
- Do NOT create a criterion for document format, page count, IEEE template
  compliance, grammar, or presentation quality. A separate fixed criterion worth
  {IEEE_POINTS} points already covers IEEE format, and it will be added to your
  criteria automatically. Your {content_points} points are for CONTENT only.

{GRADING_POLICY}

Respond with JSON in exactly this shape:
{{
  "track": {track},
  "track_name": "{TRACKS[track]}",
  "criteria": [
    {{
      "criterion": "short name",
      "max_points": integer,
      "description": "what a submission must show to earn full credit on this criterion"
    }}
  ]
}}

PROJECT INSTRUCTION DOCUMENT:
{self.instructions}
"""
        raw = self.client.call(system, user, MAX_TOKENS_RUBRIC)
        rubric = parse_json(raw)

        # Defensively drop any format criterion the model produced anyway, so it
        # cannot be double-counted against the fixed one.
        criteria = [
            c
            for c in rubric["criteria"]
            if not re.search(r"ieee|format|template|presentation", str(c["criterion"]), re.I)
        ]

        total = sum(int(c["max_points"]) for c in criteria)
        if total != content_points:
            print(
                f"  WARNING: Track {track} content criteria summed to {total}, "
                f"rescaling to {content_points}"
            )
            criteria = rescale_criteria(criteria, content_points)

        # Fixed IEEE criterion always leads the rubric.
        rubric["criteria"] = [dict(IEEE_CRITERION)] + criteria
        return rubric

    def load_or_build_rubrics(self, cache_path: Path, refresh: bool = False) -> None:
        """Build both track rubrics, caching them so runs stay consistent."""
        if cache_path.exists() and not refresh:
            cached = json.loads(cache_path.read_text(encoding="utf-8"))
            if cached.get("model") == self.client.model:
                self.rubrics = {int(k): v for k, v in cached["rubrics"].items()}
                print(f"Loaded cached rubrics from {cache_path.name}")
                return
            print("Cached rubrics were built with a different model - regenerating.")

        for track in (1, 2):
            print(f"Generating Track {track} rubric from the Project Instruction...")
            self.rubrics[track] = self.build_rubric(track)

        cache_path.write_text(
            json.dumps(
                {"model": self.client.model, "total_points": TOTAL_POINTS,
                 "rubrics": {str(k): v for k, v in self.rubrics.items()}},
                indent=2,
            ),
            encoding="utf-8",
        )
        print(f"Rubrics saved to {cache_path.name}")

    # --- step 3: classify track ---------------------------------------------

    def classify_track(self, report_text: str) -> Dict[str, Any]:
        system = (
            "You classify student project reports into one of two tracks. "
            "Always respond with valid JSON only."
        )
        user = f"""
A student team chose ONE of these two project tracks:

Track 1 - {TRACKS[1]}:
  The team builds something. They define a problem, obtain a dataset, implement
  and train an actual model or system, run experiments, and report empirical
  results (accuracy, F1, RMSE, confusion matrices, training curves, etc.).

Track 2 - {TRACKS[2]}:
  The team does NOT build a working system. They review 8-12 existing papers,
  compare approaches across that literature, and then PROPOSE a system design
  (architecture, model selection, data requirements, evaluation metrics) that is
  justified by the literature but not implemented and empirically evaluated.

The decisive question: did the team train/run a model on data and report their own
experimental results (Track 1), or did they synthesize literature and propose a
design without building it (Track 2)?

Respond with JSON:
{{
  "track": 1 or 2,
  "confidence": "high" | "medium" | "low",
  "rationale": "one or two sentences citing what in the report decided it"
}}

REPORT:
{report_text[:MAX_REPORT_CHARS]}
"""
        raw = self.client.call(system, user, MAX_TOKENS_CLASSIFY)
        result = parse_json(raw)
        result["track"] = int(result["track"])
        if result["track"] not in TRACKS:
            raise ValueError(f"Model returned an invalid track: {result['track']}")
        return result

    # --- step 4: grade criterion by criterion --------------------------------

    def grade(self, report_text: str, track: int) -> Dict[str, Any]:
        rubric = self.rubrics[track]
        criteria_block = "\n".join(
            f"{i}. {c['criterion']} (max {c['max_points']} points)\n   {c['description']}"
            for i, c in enumerate(rubric["criteria"], start=1)
        )
        system = (
            "You are an experienced but VERY generous university professor grading "
            "graduate ML project reports. You default to full credit. You never "
            "deduct without quoting specific evidence. Always respond with valid JSON only."
        )
        user = f"""
Grade this student report against the Track {track} rubric below.

TRACK: {track} - {TRACKS[track]}

RUBRIC (total {TOTAL_POINTS} points):
{criteria_block}

{GRADING_POLICY}

NOTE ON THE IEEE FORMAT CRITERION: you are reading text extracted from a PDF or
DOCX, so column layout, fonts, and margins are NOT visible to you. Judge IEEE
format only on structural markers you can actually observe in the text (abstract,
index terms, numbered sections, figure/table captions, numbered reference list).
Never deduct for layout you cannot verify. If in doubt, award full credit.

For EVERY criterion you must provide:
- "score": integer, 0 to that criterion's max_points. Default to max_points.
- "evidence": a short direct quote or specific reference from the report showing
  where the requirement is addressed. This is REQUIRED even when you give full credit.
- "deduction_reason": "" (empty string) if full credit. Otherwise the specific,
  substantive rubric violation, citing evidence from the report. Never write a
  deduction reason based on formatting, style, grammar, or wanting more depth.
- "feedback": one or two brief, constructive sentences for the students.

Respond with JSON in exactly this shape:
{{
  "criteria": [
    {{
      "criterion": "must exactly match the rubric criterion name",
      "max_points": integer,
      "score": integer,
      "evidence": "...",
      "deduction_reason": "",
      "feedback": "..."
    }}
  ],
  "overall_feedback": "3-5 sentences summarizing the report's strengths and any genuinely substantive gaps"
}}

REPORT:
{report_text[:MAX_REPORT_CHARS]}
"""
        raw = self.client.call(system, user, MAX_TOKENS_GRADE)
        result = parse_json(raw)

        # Reconcile the model's criteria against the authoritative rubric:
        # clamp scores, fill in anything the model skipped, drop inventions.
        by_name = {c.get("criterion", ""): c for c in result.get("criteria", [])}
        reconciled = []
        for spec in rubric["criteria"]:
            name = spec["criterion"]
            max_points = int(spec["max_points"])
            got = by_name.get(name) or next(
                (v for k, v in by_name.items() if k.lower().strip() == name.lower().strip()),
                None,
            )
            if got is None:
                # Model omitted this criterion. Under a lenient policy an
                # un-assessed criterion earns full credit, not zero.
                reconciled.append(
                    {
                        "criterion": name,
                        "max_points": max_points,
                        "score": max_points,
                        "evidence": "",
                        "deduction_reason": "",
                        "feedback": "Not separately assessed; full credit applied.",
                    }
                )
                continue
            score = int(round(float(got.get("score", max_points))))
            score = max(0, min(score, max_points))
            reconciled.append(
                {
                    "criterion": name,
                    "max_points": max_points,
                    "score": score,
                    "evidence": str(got.get("evidence", "")).strip(),
                    "deduction_reason": str(got.get("deduction_reason", "")).strip(),
                    "feedback": str(got.get("feedback", "")).strip(),
                }
            )

        result["criteria"] = reconciled
        # Total is computed from the criterion scores, never taken from the model.
        result["total_score"] = sum(c["score"] for c in reconciled)
        result["overall_feedback"] = str(result.get("overall_feedback", "")).strip()
        return result

    # --- orchestration -------------------------------------------------------

    def grade_group(self, label: str, report_path: Path) -> Dict[str, Any]:
        text = extract_text(report_path)
        if not text.strip():
            return {
                "group": label,
                "file": report_path.name,
                "error": "No text could be extracted from the report",
            }

        print(f"  Extracted {len(text):,} characters from {report_path.name}")

        classification = self.classify_track(text)
        track = classification["track"]
        print(
            f"  Track {track} ({classification.get('confidence', '?')} confidence): "
            f"{classification.get('rationale', '')}"
        )

        graded = self.grade(text, track)
        print(f"  Score: {graded['total_score']}/{TOTAL_POINTS}")

        return {
            "group": label,
            "file": report_path.name,
            "track": track,
            "track_name": TRACKS[track],
            "track_confidence": classification.get("confidence", ""),
            "track_rationale": classification.get("rationale", ""),
            "criteria": graded["criteria"],
            "total_score": graded["total_score"],
            "overall_feedback": graded["overall_feedback"],
        }

    def run(self, root: Path, limit: int = 0,
            only: Optional[List[int]] = None) -> List[Dict[str, Any]]:
        groups = discover_groups(root)
        print(f"\nFound {len(groups)} group folders in {root}")
        if only:
            wanted = set(only)
            groups = [
                g for g in groups
                if (m := re.search(r"Group\s*(\d+)", g[0], re.I)) and int(m.group(1)) in wanted
            ]
            print(f"Grading only group(s): {', '.join(str(n) for n in sorted(wanted))}")
        if limit:
            groups = groups[:limit]
            print(f"Limiting this run to the first {len(groups)} group(s)")
        print()

        results = []
        for i, (label, _folder, report_path) in enumerate(groups, start=1):
            print(f"[{i}/{len(groups)}] {label}")
            if report_path is None:
                print("  ERROR: no PDF/DOCX report found in this folder")
                results.append(
                    {"group": label, "file": "", "error": "No report file found"}
                )
                continue
            try:
                results.append(self.grade_group(label, report_path))
            except Exception as e:
                print(f"  ERROR: {e}")
                results.append(
                    {"group": label, "file": report_path.name, "error": str(e)}
                )
            print()
        return results


def rescale_criteria(criteria: List[Dict[str, Any]], target: int) -> List[Dict[str, Any]]:
    """Force criterion max_points to sum to target, preserving relative weights."""
    current = sum(int(c["max_points"]) for c in criteria)
    if current <= 0:
        base, extra = divmod(target, len(criteria))
        for i, c in enumerate(criteria):
            c["max_points"] = base + (1 if i < extra else 0)
        return criteria
    for c in criteria:
        c["max_points"] = max(1, round(int(c["max_points"]) * target / current))
    # Fix any residual drift from rounding on the largest criterion.
    drift = target - sum(c["max_points"] for c in criteria)
    if drift:
        biggest = max(criteria, key=lambda c: c["max_points"])
        biggest["max_points"] = max(1, biggest["max_points"] + drift)
    return criteria


# ============================================================================
# OUTPUT
# ============================================================================


def write_csv(results: List[Dict[str, Any]], path: Path) -> None:
    """One row per criterion per group - the structured, machine-readable output."""
    with open(path, "w", newline="", encoding="utf-8") as f:
        writer = csv.writer(f)
        writer.writerow(
            [
                "Group", "File", "Track", "Track Name", "Track Confidence",
                "Criterion", "Score", "Max Points", "Evidence",
                "Deduction Reason", "Feedback",
                "Group Total", "Group Max", "Group Percentage", "Overall Feedback",
            ]
        )
        for r in results:
            if "error" in r:
                writer.writerow(
                    [r["group"], r.get("file", ""), "ERROR", "", "", "", "", "",
                     "", r["error"], "", "", TOTAL_POINTS, "", ""]
                )
                continue
            pct = round(r["total_score"] / TOTAL_POINTS * 100)
            for c in r["criteria"]:
                writer.writerow(
                    [
                        r["group"], r["file"], r["track"], r["track_name"],
                        r["track_confidence"], c["criterion"], c["score"],
                        c["max_points"], c["evidence"], c["deduction_reason"],
                        c["feedback"], r["total_score"], TOTAL_POINTS,
                        f"{pct}%", r["overall_feedback"],
                    ]
                )


def write_summary_csv(results: List[Dict[str, Any]], path: Path) -> None:
    """One row per group - the at-a-glance gradebook."""
    rows = []
    for r in results:
        if "error" in r:
            rows.append({"Group": r["group"], "File": r.get("file", ""),
                         "Track": "ERROR", "Total": "", "Max": TOTAL_POINTS,
                         "Percentage": "", "Error": r["error"]})
            continue
        rows.append(
            {
                "Group": r["group"],
                "File": r["file"],
                "Track": f"Track {r['track']}",
                "Total": r["total_score"],
                "Max": TOTAL_POINTS,
                "Percentage": f"{round(r['total_score'] / TOTAL_POINTS * 100)}%",
                "Error": "",
            }
        )
    pd.DataFrame(rows).to_csv(path, index=False, encoding="utf-8")


def write_excel(results: List[Dict[str, Any]], path: Path) -> None:
    summary, detail = [], []
    for r in results:
        if "error" in r:
            summary.append({"Group": r["group"], "Track": "ERROR",
                            "Total Score": "", "Percentage": "", "Note": r["error"]})
            continue
        pct = round(r["total_score"] / TOTAL_POINTS * 100)
        summary.append(
            {
                "Group": r["group"],
                "Track": f"Track {r['track']} - {r['track_name']}",
                "Total Score": f"{r['total_score']}/{TOTAL_POINTS}",
                "Percentage": f"{pct}%",
                "Note": "",
            }
        )
        for c in r["criteria"]:
            detail.append(
                {
                    "Group": r["group"], "Track": r["track"],
                    "Criterion": c["criterion"],
                    "Score": f"{c['score']}/{c['max_points']}",
                    "Evidence": c["evidence"],
                    "Deduction Reason": c["deduction_reason"],
                    "Feedback": c["feedback"],
                }
            )

    with pd.ExcelWriter(path, engine="openpyxl") as writer:
        for name, rows in (("Summary", summary), ("Criterion Detail", detail)):
            if not rows:
                continue
            df = pd.DataFrame(rows)
            df.to_excel(writer, index=False, sheet_name=name)
            ws = writer.sheets[name]
            for idx, column in enumerate(df.columns, start=1):
                width = max(df[column].astype(str).map(len).max(), len(column))
                ws.column_dimensions[get_column_letter(idx)].width = min(width + 2, 60)


def write_markdown(results: List[Dict[str, Any]], rubrics: Dict[int, Dict[str, Any]],
                   path: Path) -> None:
    """The readable grading report."""
    lines: List[str] = []
    lines.append("# MIA 5100 - Project Report Grading")
    lines.append("")
    lines.append(f"Generated: {datetime.now():%Y-%m-%d %H:%M}")
    lines.append(f"Maximum score: {TOTAL_POINTS} points per report")
    lines.append("")

    # Gradebook
    lines.append("## Summary")
    lines.append("")
    lines.append("| Group | Track | Score | Percentage |")
    lines.append("|---|---|---|---|")
    for r in results:
        if "error" in r:
            lines.append(f"| {r['group']} | ERROR | - | - |")
            continue
        pct = round(r["total_score"] / TOTAL_POINTS * 100)
        lines.append(
            f"| {r['group']} | Track {r['track']} | "
            f"{r['total_score']}/{TOTAL_POINTS} | {pct}% |"
        )
    lines.append("")

    # Rubrics used
    for track in sorted(rubrics):
        rubric = rubrics[track]
        lines.append(f"## Rubric - Track {track}: {TRACKS[track]}")
        lines.append("")
        lines.append("| Criterion | Max | What earns full credit |")
        lines.append("|---|---|---|")
        for c in rubric["criteria"]:
            desc = str(c["description"]).replace("|", "\\|")
            lines.append(f"| {c['criterion']} | {c['max_points']} | {desc} |")
        lines.append("")

    # Per-group detail
    lines.append("## Group Reports")
    lines.append("")
    for r in results:
        lines.append(f"### {r['group']}")
        lines.append("")
        if "error" in r:
            lines.append(f"**ERROR:** {r['error']}")
            lines.append("")
            continue
        pct = round(r["total_score"] / TOTAL_POINTS * 100)
        lines.append(f"- **File:** {r['file']}")
        lines.append(f"- **Track:** {r['track']} - {r['track_name']}")
        lines.append(
            f"- **Track determination** ({r['track_confidence']} confidence): "
            f"{r['track_rationale']}"
        )
        lines.append(f"- **Total score:** {r['total_score']}/{TOTAL_POINTS} ({pct}%)")
        lines.append("")
        for c in r["criteria"]:
            lines.append(f"#### {c['criterion']} - {c['score']}/{c['max_points']}")
            lines.append("")
            if c["evidence"]:
                lines.append(f"- *Evidence:* {c['evidence']}")
            if c["deduction_reason"]:
                lines.append(f"- *Deduction:* {c['deduction_reason']}")
            else:
                lines.append("- *Deduction:* none - full credit")
            if c["feedback"]:
                lines.append(f"- *Feedback:* {c['feedback']}")
            lines.append("")
        lines.append(f"**Overall feedback:** {r['overall_feedback']}")
        lines.append("")
        lines.append("---")
        lines.append("")

    path.write_text("\n".join(lines), encoding="utf-8")


def print_console_summary(results: List[Dict[str, Any]]) -> None:
    print("=" * 78)
    print("GRADING SUMMARY")
    print("=" * 78)
    scored = []
    for r in results:
        if "error" in r:
            print(f"{r['group']:<12} ERROR: {r['error']}")
            continue
        pct = round(r["total_score"] / TOTAL_POINTS * 100)
        scored.append(r["total_score"])
        print(
            f"{r['group']:<12} Track {r['track']}   "
            f"{r['total_score']:>2}/{TOTAL_POINTS}  ({pct}%)"
        )
    if scored:
        print("-" * 78)
        print(
            f"{'Mean':<12}         {sum(scored) / len(scored):.1f}/{TOTAL_POINTS}   "
            f"min {min(scored)}   max {max(scored)}   n={len(scored)}"
        )
    print("=" * 78)


# ============================================================================
# MAIN
# ============================================================================


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Grade MIA 5100 project reports")
    p.add_argument("--root", default=ROOT_FOLDER, help="Root folder of group subfolders")
    p.add_argument("--instructions", default=INSTRUCTIONS_FILE, help="Project Instruction file")
    p.add_argument("--model", default=MODEL, help="Gateway model name")
    p.add_argument("--limit", type=int, default=0, help="Only grade the first N groups")
    p.add_argument("--groups", default="", help="Only grade these group numbers, e.g. 6,12")
    p.add_argument("--refresh-rubrics", action="store_true", help="Rebuild cached rubrics")
    p.add_argument("--outdir", default=".", help="Where to write the output files")
    return p.parse_args()


def main() -> None:
    args = parse_args()

    if not API_KEY or API_KEY == "YOUR_API_KEY_HERE":
        print("Please set GENAI_API_KEY in the .env file (Geotab GenAI Gateway token).")
        return

    root = Path(args.root)
    if not root.is_dir():
        print(f"Error: root folder not found: {root}")
        return

    instructions_path = Path(args.instructions)
    if not instructions_path.is_file():
        print(f"Error: instruction document not found: {instructions_path}")
        return

    print(f"Gateway:      {BASE_URL}")
    print(f"Model:        {args.model}")
    print(f"Instructions: {instructions_path.name}")
    print(f"Root:         {root}")
    print()

    instructions_text = extract_text(instructions_path)
    if not instructions_text.strip():
        print("Error: could not extract text from the instruction document.")
        return
    print(f"Read {len(instructions_text):,} characters of project instructions.")

    client = GatewayClient(API_KEY, BASE_URL, args.model)
    grader = ProjectGrader(client, instructions_text)
    grader.load_or_build_rubrics(Path("rubrics.json"), refresh=args.refresh_rubrics)

    for track in sorted(grader.rubrics):
        criteria = grader.rubrics[track]["criteria"]
        total = sum(c["max_points"] for c in criteria)
        print(f"  Track {track}: {len(criteria)} criteria, {total} points")
    print()

    only = [int(x) for x in re.findall(r"\d+", args.groups)] or None
    results = grader.run(root, limit=args.limit, only=only)

    outdir = Path(args.outdir)
    outdir.mkdir(parents=True, exist_ok=True)
    stamp = datetime.now().strftime("%Y%m%d_%H%M%S")

    detail_csv = outdir / f"grading_detail_{stamp}.csv"
    summary_csv = outdir / f"grading_summary_{stamp}.csv"
    excel = outdir / f"grading_results_{stamp}.xlsx"
    markdown = outdir / f"grading_report_{stamp}.md"
    raw_json = outdir / f"grading_raw_{stamp}.json"

    write_csv(results, detail_csv)
    write_summary_csv(results, summary_csv)
    write_excel(results, excel)
    write_markdown(results, grader.rubrics, markdown)
    raw_json.write_text(json.dumps(results, indent=2), encoding="utf-8")

    print()
    print_console_summary(results)
    print()
    print("Output files:")
    for f in (summary_csv, detail_csv, excel, markdown, raw_json):
        print(f"  {f}")


if __name__ == "__main__":
    main()
